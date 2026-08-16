#!/usr/bin/env python3
"""Build the privacy-safe AIMUNC 2026 graduate-review project brief."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "artifacts" / "AIMUNC2026_项目说明书_申研版.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "0B2545"
GRAY = "5B6573"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"
GOLD = "A66A00"
BLACK = "111111"


def set_run_font(run, size=None, *, bold=None, italic=None, color=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_fill(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, *, indent=120, borders=True):
    if sum(widths) != 9360:
        raise ValueError("Table widths must total 9360 DXA")

    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    border_node = tbl_pr.find(qn("w:tblBorders"))
    if border_node is None:
        border_node = OxmlElement("w:tblBorders")
        tbl_pr.append(border_node)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = border_node.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            border_node.append(node)
        node.set(qn("w:val"), "single" if borders else "nil")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:color"), "D9DEE5")

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, separate, text, end))
    set_run_font(run, 9, color=GRAY)


def add_numbering(doc):
    numbering = doc.part.numbering_part.element
    existing_abstract = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    abstract_id = max(existing_abstract, default=-1) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•")
    level.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    level.append(suff)

    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)

    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r_pr.append(fonts)
    level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    existing_num = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    num_id = max(existing_num, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_bullet(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend((ilvl, num))
    p_pr.append(num_pr)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.167


def add_body(doc, text, *, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    p.paragraph_format.widow_control = True
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, 11, bold=True, color=BLACK)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest, 11, color=BLACK)
    else:
        run = p.add_run(text)
        set_run_font(run, 11, color=BLACK)
    return p


def add_bullet(doc, text, num_id, *, bold_lead=None):
    p = doc.add_paragraph()
    apply_bullet(p, num_id)
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, 11, bold=True, color=BLACK)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest, 11, color=BLACK)
    else:
        run = p.add_run(text)
        set_run_font(run, 11, color=BLACK)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    keep_with_next(p)
    return p


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_fill(cell, CALLOUT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.10
    lead = p.add_run(title + "  ")
    set_run_font(lead, 10.5, bold=True, color=DARK_BLUE)
    body = p.add_run(text)
    set_run_font(body, 10.5, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    tokens = {
        1: (16, BLUE, 16, 8),
        2: (13, BLUE, 12, 6),
        3: (12, DARK_BLUE, 8, 4),
    }
    for level, (size, color, before, after) in tokens.items():
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True


def configure_section(section):
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    p = header.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    left = p.add_run("AIMUNC 2026 · 项目说明书")
    set_run_font(left, 8.5, bold=True, color=GRAY)
    p.add_run("\t")
    right = p.add_run("申研评审材料｜脱敏版")
    set_run_font(right, 8.5, color=GRAY)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    prefix = fp.add_run("第 ")
    set_run_font(prefix, 9, color=GRAY)
    add_field(fp, "PAGE")
    suffix = fp.add_run(" 页")
    set_run_font(suffix, 9, color=GRAY)


def add_metadata_table(doc):
    table = doc.add_table(rows=4, cols=4)
    set_table_geometry(table, [1200, 3480, 1200, 3480])
    rows = [
        ("项目类型", "独立全栈 / 产品工程", "开发周期", "2026-01-19 至 2026-03-17"),
        ("设计规模", "约 250–300 人会议", "当前状态", "实际服务 237 名代表"),
        ("角色体系", "代表 / 领队 / 管理员", "技术栈", "Java + Spring Boot + MySQL + 原生 Web"),
        ("个人角色", "唯一人类开发者", "AI 使用", "Claude 辅助研发；产品无 Agent"),
    ]
    for row_index, values in enumerate(rows):
        row = table.rows[row_index]
        for col_index, value in enumerate(values):
            cell = row.cells[col_index]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(value)
            if col_index % 2 == 0:
                set_cell_fill(cell, LIGHT_GRAY)
                set_run_font(run, 9.5, bold=True, color=DARK_BLUE)
            else:
                set_run_font(run, 9.5, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_role_table(doc):
    table = doc.add_table(rows=4, cols=3)
    set_table_geometry(table, [1700, 3830, 3830])
    headers = ("角色", "主要任务", "系统支持")
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_fill(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.paragraph_format.keep_with_next = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(text)
        set_run_font(run, 9.5, bold=True, color=DARK_BLUE)
    set_repeat_table_header(table.rows[0])

    rows = (
        ("代表", "提交信息、志愿和付款凭证；查询状态", "个人报名页、状态页、受控上传与查询"),
        ("领队", "创建代表团、邀请成员、执行初审", "邀请码、成员列表、审核与团队进度"),
        ("管理员", "查看全局数据、终审、分配会场、审核付款", "双视图、统计、筛选排序、分页与导出"),
    )
    for row_index, values in enumerate(rows, start=1):
        for col_index, value in enumerate(values):
            cell = table.rows[row_index].cells[col_index]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.05
            if col_index == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(value)
            set_run_font(run, 9.2, bold=(col_index == 0), color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_evidence_table(doc):
    table = doc.add_table(rows=6, cols=3)
    set_table_geometry(table, [2500, 1800, 5060])
    headers = ("主张", "状态", "安全表述")
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_fill(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(text)
        set_run_font(run, 9.2, bold=True, color=DARK_BLUE)
    set_repeat_table_header(table.rows[0])

    rows = (
        ("250–300 人", "设计容量", "面向约 250–300 人规模会议设计"),
        ("237 名代表", "用户确认", "正式使用阶段实际服务；待补脱敏聚合证据"),
        ("多端同步", "部分成立", "多角色端数据联动；请求刷新后保持一致"),
        ("AI / Agent", "研发侧 AI", "Claude 辅助研发；平台无产品内 Agent"),
        ("测试", "覆盖有限", "手工 API/E2E 联调 + 1 个上下文 smoke test"),
    )
    for row_index, values in enumerate(rows, start=1):
        for col_index, value in enumerate(values):
            cell = table.rows[row_index].cells[col_index]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.0
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_index == 1 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(value)
            color = GOLD if value in {"未核验", "部分成立", "覆盖有限"} else BLACK
            set_run_font(run, 9.0, bold=(col_index == 1), color=color)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def build():
    doc = Document()
    configure_styles(doc)
    for section in doc.sections:
        configure_section(section)
    num_id = add_numbering(doc)

    props = doc.core_properties
    props.title = "AIMUNC 2026 项目说明书（申研版）"
    props.subject = "独立全栈项目、产品工程与 AI 辅助开发复盘"
    props.author = ""
    props.last_modified_by = ""
    props.keywords = "AIMUNC, full-stack, product engineering, AI-assisted development"

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(6)
    kicker.paragraph_format.space_after = Pt(6)
    run = kicker.add_run("PROJECT PORTFOLIO")
    set_run_font(run, 10, bold=True, color=BLUE)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(5)
    keep_with_next(title)
    run = title.add_run("AIMUNC 2026")
    set_run_font(run, 28, bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(5)
    keep_with_next(subtitle)
    run = subtitle.add_run("模拟联合国会议信息收集、报名与管理平台")
    set_run_font(run, 15, bold=True, color=DARK_BLUE)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_before = Pt(0)
    meta.paragraph_format.space_after = Pt(15)
    run = meta.add_run("申研评审材料 · 脱敏版 · 2026-08-16")
    set_run_font(run, 9.5, color=GRAY)

    add_metadata_table(doc)
    add_callout(
        doc,
        "事实边界",
        "系统按约 250–300 人会议规模设计；早期版本经过小范围测试，后续正式使用阶段实际服务 237 名代表。237 为项目负责人确认的使用口径，待补脱敏聚合证据，不代表并发或性能指标；文档与仓库均不包含真实报名数据。",
    )

    add_heading(doc, "1. 项目概述", 1)
    add_body(
        doc,
        "AIMUNC 2026 将分散在表格和聊天工具中的报名、代表团关系、志愿、审核状态和付款凭证整合到统一平台。代表、领队和管理员通过同一后端与数据库协作，减少人工转录，并为会务组织者提供可操作的全局视图。",
    )
    add_body(
        doc,
        "这是我的第一个独立全栈项目。除业务规则和运营需求与社团讨论外，我是唯一的人类开发者，独立承担产品梳理、数据建模、前后端实现、管理端设计、测试、部署和迭代。Claude 用于架构讨论、编码辅助和审计；平台本身没有内置 Agent 功能。",
    )

    section_two = add_heading(doc, "2. 问题与产品设计", 1)
    section_two.paragraph_format.page_break_before = True
    add_body(
        doc,
        "大型学生会议的难点不只是收集表单，而是让不同角色围绕同一份状态持续工作。我先将需求拆成三类用户旅程，再定义用户、代表团和代表报名三类核心实体，以及报名、初审、终审、会场分配和付款审核的状态流。",
    )
    add_role_table(doc)
    add_body(
        doc,
        "数据联动方式：代表提交后，领队在成员列表中完成初审，管理员在全局后台完成终审与分配；代表再次加载状态页时得到最新结果。这里的“多端联动”依赖共享数据库和 REST 请求刷新，并非 WebSocket 实时推送。",
    )

    add_heading(doc, "3. 我的职责与工程实现", 1)
    responsibilities = (
        ("需求转译：", "与社团讨论会务流程，将非结构化运营要求转化为页面、字段、角色和状态约束。"),
        ("全栈实现：", "使用 Java 17、Spring Boot、JPA、MySQL 与原生 Web 技术完成后端、数据库和 9 个业务页面。"),
        ("管理端设计：", "实现代表/代表团双视图、统计看板、多维筛选排序、分页、终审、会场分配、付款审核和 Excel 导出。"),
        ("安全演进：", "引入 BCrypt、JWT、角色与资源归属校验，并在展示仓库重建时继续修补已发现的权限缺口。"),
        ("测试与部署：", "完成手工 API/E2E 联调、问题复现和回归，使用 VPS、Cloudflare、Nginx 与 systemd 部署。"),
    )
    for lead, body in responsibilities:
        add_bullet(doc, lead + body, num_id, bold_lead=lead)

    add_heading(doc, "4. 迭代与结果", 1)
    add_body(
        doc,
        "项目于 2026-01-19 至 2026-03-17 从 v0.1 迭代至 v3.5。当前快照约含 15 个主要 Java 文件、9 个业务页面、20 个 REST 接口和约 1.16 万行文本；该口径包含 Java、HTML、CSS、JavaScript、注释和页面样式，不等同于纯业务代码行。",
    )
    add_body(
        doc,
        "完成服务器部署后，早期版本通过小范围测试记录 8 项反馈，涉及手机端阅读、付款凭证引导、主题一致性和权限问题；后续正式使用阶段实际服务 237 名代表。日志保留了优先级排序、跨文件任务清单、回退和复测过程。",
    )
    add_callout(
        doc,
        "测试口径",
        "当前自动化测试只有 1 个 Spring 上下文 smoke test；更广的验证来自手工 API 调试、端到端流程和回归清单。因此不宣称拥有完整自动化测试体系。",
    )

    add_heading(doc, "5. AI 辅助开发：真实作用与责任边界", 1)
    add_body(
        doc,
        "Claude Opus/Sonnet 在研发过程中曾充当架构讨论者和代码助手，参与任务拆解、实现建议、错误定位、跨文件审计和开发日志交接。我负责决定需求是否进入产品、选择或拒绝方案、整合代码、运行验证、部署并承担结果。",
    )
    add_body(
        doc,
        "日志也记录了 AI 修改造成页面回归后由人工发现、回退和修复的过程。这让我认识到，AI 可以降低进入陌生技术领域的门槛，却不能替代开发者对上下文、权限、验证和交付质量的判断。更准确的描述是“唯一人类开发者 + AI 辅助研发”，而不是多人开发团队，也不是产品内 Agent。",
    )

    add_heading(doc, "6. 已知不足与工程反思", 1)
    limitations = (
        "原生前端页面较大，组件、状态与样式复用不足。",
        "部分 Controller 同时承担校验、业务逻辑、DTO 拼装与文件操作。",
        "JWT 已引入，但授权仍主要分散在 Controller，尚未形成统一策略。",
        "数据库迁移、一键部署、监控、备份演练与系统化测试仍不完整。",
        "237 名代表实际使用不等于并发或性能验证，且仍待补脱敏聚合证据。",
        "当前产品没有 Agent 功能，不能把 AI 辅助开发等同于 AI 产品能力。",
    )
    for item in limitations:
        add_bullet(doc, item, num_id)
    add_body(
        doc,
        "我选择保留这些不足，因为第一版的价值不是证明它已经完美，而是证明我能够把一个实际组织问题推进到可运行、可部署、可反馈迭代的系统，并在事后识别结构性问题。",
    )

    add_heading(doc, "7. Python / Agent V2 路线", 1)
    add_body(
        doc,
        "后续将保留 Java/Spring v1 作为历史基线，以独立 V2 验证技术成长，而不是覆盖第一版。建议采用 FastAPI、Pydantic、SQLAlchemy、Alembic 和 pytest 构建模块化单体，补齐类型化接口、数据库迁移、权限测试、结构化日志和可复现部署。",
    )
    v2_items = (
        "报名完整性检查、缺失字段和异常状态提示。",
        "仅对聚合或脱敏数据开放的自然语言进度查询。",
        "通知草稿与反馈汇总，所有外部动作必须人工确认。",
        "会场分配建议，但保留规则解释、人工审批和操作日志。",
        "用任务成功率、事实准确率、越权率和人工纠错成本评估，而不是以“接入模型”作为完成标准。",
    )
    for item in v2_items:
        add_bullet(doc, item, num_id)
    add_callout(
        doc,
        "规划声明",
        "上述 Python 与 Agent 内容均为后续方向，不是当前版本已经实现的成果。Agent 不应直接访问或自由输出参与者个人信息，也不应自动执行录取、终审等高影响决定。",
    )

    add_heading(doc, "8. 与个人发展方向的关系", 1)
    add_body(
        doc,
        "AIMUNC 2026 让我第一次完整经历“理解真实需求—构建产品—完成工程交付—反思并规划重构”的过程。它连接了我希望继续发展的三个方向：AI 产品经理需要把复杂运营问题转化为清晰产品边界；Agent 应用开发需要寻找可验证、可控的真实任务；FDE 需要深入业务现场，与使用者协作并独立推动方案落地。",
    )
    add_body(
        doc,
        "对我而言，这不是一个为了展示新技术而包装的项目，而是我建立产品与工程能力的起点。下一阶段的重点不是掩盖第一版的不足，而是以可复现、可评估的方法证明自己如何从这些不足中成长。",
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
